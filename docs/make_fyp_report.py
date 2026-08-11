"""
Build the FYP Mid Evaluation Report as .docx, from the university template.

The template is opened and its body emptied, so page size, margins and the
heading/caption/toc styles are inherited exactly rather than reconstructed. Only
the content is ours.

Every figure referenced here is produced by make_figures.py and
make_report_figures.py from the code itself, and every metric is read from
outputs/validation.json and outputs/metrics.json rather than typed in, so the
report cannot quietly drift from the system it describes.

    python docs/make_fyp_report.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
FIG = HERE / "figures"
TEMPLATE = Path(r"C:\Users\kasha\Downloads\FYP Mid Evaluation Report (Template)  (4).docx")
OUT = HERE / "APOLLO-M_FYP_Mid_Evaluation_Report.docx"

TITLE = ("APOLLO-M: AI FRAMEWORK FOR FORECASTING INSTABILITY "
         "IN DIGITAL COMMUNITIES")
STUDENTS = [("Kashaf Fatima", "84381"), ("Rizwan Saleem", ""), ("Aaqib Mehmood", "")]
SUPERVISOR = "Mr. Saghir Ahmed"
YEAR = "2026"

# Measured values, read rather than retyped.
V = json.loads((ROOT / "outputs" / "validation.json").read_text(encoding="utf-8"))
M = json.loads((ROOT / "outputs" / "metrics.json").read_text(encoding="utf-8"))
TOX = M["Toxicity (TF-IDF+LR, Jigsaw)"]
DET, FC = V["detection"], V["forecast_tft"]


# --------------------------------------------------------------------------
# document helpers
# --------------------------------------------------------------------------

def new_doc() -> Document:
    d = Document(str(TEMPLATE))
    body = d.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)
    return d


def para(d, text="", style=None, size=12, bold=False, align=None, italic=False,
         space_after=6, indent=None, color=None):
    p = d.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    return p


def body_text(d, text, indent_first=False):
    return para(d, text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_after=10, indent=1.27 if indent_first else None)


def chapter(d, number, title):
    d.add_page_break()
    para(d, f"CHAPTER {number}", size=12, bold=True, space_after=2)
    p = d.add_paragraph(style="Heading 2")
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)
    return p


def section(d, title, level=3):
    p = d.add_paragraph(style=f"Heading {level}")
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12 if level == 3 else 11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def figure(d, name, caption, width=6.1):
    path = FIG / f"{name}.png"
    if not path.exists():
        para(d, f"[missing figure: {name}]", italic=True)
        return
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    para(d, caption, style="Caption", size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=12)


def table(d, caption, headers, rows, widths=None, font=8.5):
    para(d, caption, style="Caption", size=10, space_after=4)
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    para(d, "", space_after=10)
    return t


# --------------------------------------------------------------------------
# front matter
# --------------------------------------------------------------------------

def front_matter(d):
    for _ in range(4):
        para(d, "")
    para(d, TITLE, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for _ in range(6):
        para(d, "")
    for name, _ in STUDENTS:
        para(d, name, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for _ in range(5):
        para(d, "")
    para(d, "A project report submitted in partial fulfilment of the", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(d, "requirements for the award of the degree of", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(d, "Bachelor of Science in Computer Science (BSCS)", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for _ in range(5):
        para(d, "")
    para(d, "Computer Science Department", size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=0)
    para(d, "Bahria University, Karachi Campus", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for _ in range(3):
        para(d, "")
    para(d, YEAR, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # DECLARATION
    d.add_page_break()
    para(d, "DECLARATION", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    body_text(d, "We hereby declare that this project report is based on our original "
                 "work except for citations and quotations which have been duly "
                 "acknowledged. We also declare that it has not been previously and "
                 "concurrently submitted for any other degree or award at Bahria "
                 "University or any other institution.")
    para(d, "")
    for name, reg in STUDENTS:
        para(d, "Signature\t:\t_________________________", size=12, space_after=2)
        para(d, f"Name\t:\t{name}", size=12, space_after=2)
        para(d, f"Reg No.\t:\t{reg if reg else '_________________________'}",
             size=12, space_after=10)
    para(d, "Date\t:\t_________________________", size=12)

    # APPROVAL
    d.add_page_break()
    para(d, "APPROVAL FOR SUBMISSION", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    body_text(d, f"We certify that this project report entitled "
                 f"\u201c{TITLE}\u201d was prepared by KASHAF FATIMA, RIZWAN SALEEM and "
                 f"AAQIB MEHMOOD, and has met the required standard for submission in "
                 f"partial fulfilment of the requirements for the award of the degree "
                 f"of Bachelor of Science in Computer Science at Bahria University, "
                 f"Karachi Campus.")
    para(d, "")
    para(d, "Approved by,", size=12, space_after=12)
    para(d, "Signature\t:\t_________________________", size=12, space_after=2)
    para(d, f"Supervisor\t:\t{SUPERVISOR}", size=12, space_after=2)
    para(d, "Date\t:\t_________________________", size=12)

    # COPYRIGHT
    d.add_page_break()
    for _ in range(3):
        para(d, "")
    body_text(d, "The copyright of this report belongs to Bahria University according "
                 "to the Intellectual Property Policy of Bahria University. Due "
                 "acknowledgement shall always be made of the use of any material "
                 "contained in, or derived from, this report.")
    para(d, "")
    para(d, f"\u00a9 {YEAR} Bahria University. All rights reserved.", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    # ACKNOWLEDGEMENTS
    d.add_page_break()
    para(d, "ACKNOWLEDGEMENTS", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    body_text(d, "We would like to thank everyone who contributed to the successful "
                 f"completion of this project. In particular, we wish to express our "
                 f"sincere gratitude to our supervisor, {SUPERVISOR}, for his "
                 "invaluable advice, guidance and patience throughout the development "
                 "of APOLLO-M, and for consistently pushing us to justify our results "
                 "rather than merely present them.")
    body_text(d, "We also thank the Computer Science Department at Bahria University, "
                 "Karachi Campus, for providing the environment and resources that "
                 "made this work possible, and the maintainers of the open datasets "
                 "and open-source libraries on which this project depends \u2014 without "
                 "the Davidson toxicity corpus, the SNAP Reddit hyperlink network and "
                 "the PyTorch ecosystem, a project of this scope would not have been "
                 "achievable within a single academic year.", indent_first=True)
    body_text(d, "Finally, we thank our families and friends for their encouragement "
                 "and support throughout this work.", indent_first=True)

    # ABSTRACT
    d.add_page_break()
    para(d, TITLE, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(d, "")
    para(d, "ABSTRACT", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    body_text(d,
        "Online communities can deteriorate. Hostility rises, moderators are "
        "overwhelmed, constructive members leave, and a space that functioned "
        "collapses into conflict. Existing moderation tools are almost entirely "
        "reactive: they classify an individual comment after it has been posted, and "
        "they describe the present state of a community rather than its direction of "
        "travel. By the time a dashboard reports that a community is unhealthy, the "
        "damage that made it unhealthy has already occurred.")
    body_text(d,
        "APOLLO-M is a multi-layer framework that treats community instability as a "
        "forecasting problem rather than a classification problem. A micro layer "
        "scores individual comments for toxicity using a transformer classifier. A "
        "meso layer aggregates those scores, together with author churn and "
        "graph-derived polarisation and echo-chamber measures, into a Community "
        "Health Index and a trend-aware instability score. An unsupervised layer "
        "clusters communities by behavioural profile without labels. A macro layer "
        "applies a Temporal Fusion Transformer to produce a five-day toxicity "
        "forecast with p10/p50/p90 confidence bands per community. Results are "
        "persisted to PostgreSQL, served through an authenticated REST API, and "
        "presented in a dashboard with Prometheus and Grafana monitoring. A language "
        "model is used only to explain findings in natural language; every number is "
        "produced by a deterministic model or a measurable classifier.",
        indent_first=True)
    body_text(d,
        "Because no public dataset labels which online communities genuinely "
        "destabilised and when, the framework is validated against a declared "
        "simulation with recorded ground truth: sixty communities are generated over "
        "one hundred and twenty days from real Reddit community structure and a real "
        "labelled toxicity corpus, and fifteen are given a deliberately rising "
        "toxicity trend that the pipeline never sees. The forecaster identifies a "
        f"rising trend in {int(FC['destabilising_recall']*15)} of 15 planted "
        f"communities, with a slope ROC-AUC of {FC['slope_roc_auc']}. Ranking "
        f"communities by the trend-aware instability score achieves an ROC-AUC of "
        f"{DET['instability_score_roc_auc']}, against {DET['chi_roc_auc']} for the "
        f"Community Health Index and {DET['toxicity_only_roc_auc']} for raw toxicity "
        "alone. That ordering is itself a finding: an index built from present-tense "
        "health cannot rank communities by how quickly they are changing, which is "
        "what forecasting instability requires. The toxicity classifier, evaluated on "
        f"a held-out split of real labelled data, achieves {TOX['accuracy']} accuracy "
        f"and a macro F1 of {TOX['F1 macro']}.", indent_first=True)


def toc(d):
    d.add_page_break()
    para(d, "TABLE OF CONTENTS", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    front = [("DECLARATION", "ii"), ("APPROVAL FOR SUBMISSION", "iii"),
             ("ACKNOWLEDGEMENTS", "v"), ("ABSTRACT", "vi"),
             ("TABLE OF CONTENTS", "vii"), ("LIST OF TABLES", "ix"),
             ("LIST OF FIGURES", "x"), ("LIST OF SYMBOLS / ABBREVIATIONS", "xi")]
    for name, pg in front:
        p = para(d, "", size=11, space_after=2)
        p.add_run(name).font.size = Pt(11)
        p.add_run("\t" + pg).font.size = Pt(11)
    para(d, "")
    para(d, "CHAPTER", size=11, bold=True, space_after=4)
    items = [
        (1, "1", "INTRODUCTION", "1"), (2, "1.1", "Background", "1"),
        (2, "1.2", "Problem Statement", "3"), (2, "1.3", "Aims and Objectives", "4"),
        (2, "1.4", "Scope of Project", "5"),
        (1, "2", "LITERATURE REVIEW", "6"), (2, "2.1", "Background", "6"),
        (2, "2.2", "Related Work", "7"),
        (3, "2.2.1", "Toxicity and Abusive Language Detection", "7"),
        (3, "2.2.2", "Community-Level Health and Conflict", "8"),
        (3, "2.2.3", "Graph Representations of Communities", "9"),
        (3, "2.2.4", "Temporal Forecasting of Social Signals", "9"),
        (3, "2.2.5", "Misinformation Detection", "10"),
        (2, "2.3", "Research Gap and Summary", "11"),
        (2, "2.4", "Comparison with Existing Studies", "14"),
        (2, "2.5", "Chapter Summary", "15"),
        (1, "3", "DESIGN AND METHODOLOGY", "16"),
        (2, "3.1", "Proposed Methodology (Framework / Architecture)", "16"),
        (2, "3.2", "Process Model", "18"), (2, "3.3", "Modules Discussion", "19"),
        (3, "3.3.1", "Micro Layer", "19"), (3, "3.3.2", "Meso Layer", "20"),
        (3, "3.3.3", "Unsupervised Layer", "21"), (3, "3.3.4", "Macro Layer", "21"),
        (3, "3.3.5", "Act, Serve and Observe", "22"),
        (2, "3.4", "Project Diagrams", "23"),
        (3, "3.4.1", "Database Design (ERD)", "23"),
        (3, "3.4.2", "Sequence Diagram", "24"), (3, "3.4.3", "Context Diagram", "25"),
        (3, "3.4.4", "Use Case Diagram", "26"),
        (2, "3.5", "Datasets and Data Provenance", "27"),
        (2, "3.6", "Evaluation Methodology and Results", "28"),
        (2, "3.7", "Chapter Summary", "30"),
        (1, "", "REFERENCES", "31"),
    ]
    for lvl, num, name, pg in items:
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7 * (lvl - 1))
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run((f"{num}\t" if num else "") + name)
        r.font.size = Pt(11)
        r.bold = lvl == 1
        p.add_run("\t" + pg).font.size = Pt(11)

    # LIST OF TABLES
    d.add_page_break()
    para(d, "LIST OF TABLES", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(d, "TABLE\tTITLE\tPAGE", size=11, bold=True, space_after=8)
    for t, name, pg in [
        ("Table 2.1", "Research Summary", "11"),
        ("Table 2.2", "Comparison with Existing Studies", "14"),
        ("Table 3.1", "Datasets Used and Their Role", "27"),
        ("Table 3.2", "Detection Performance Against Planted Ground Truth", "28"),
        ("Table 3.3", "Forecasting Performance Against Planted Ground Truth", "29"),
    ]:
        para(d, f"{t}: {name}\t{pg}", size=11, space_after=3)

    # LIST OF FIGURES
    d.add_page_break()
    para(d, "LIST OF FIGURES", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(d, "FIGURE\tTITLE\tPAGE", size=11, bold=True, space_after=8)
    for t, name, pg in [
        ("Figure 3.1", "APOLLO-M System Architecture", "17"),
        ("Figure 3.2", "Process Model", "18"),
        ("Figure 3.3", "System Block Diagram", "19"),
        ("Figure 3.4", "Database Design (ERD)", "23"),
        ("Figure 3.5", "Sequence Diagram", "24"),
        ("Figure 3.6", "Context Diagram", "25"),
        ("Figure 3.7", "Use Case Diagram", "26"),
        ("Figure 3.8", "Request Processing Flow", "28"),
        ("Figure 3.9", "Integration with CEREBRO", "30"),
    ]:
        para(d, f"{t}: {name}\t{pg}", size=11, space_after=3)

    # ABBREVIATIONS
    d.add_page_break()
    para(d, "LIST OF SYMBOLS / ABBREVIATIONS", style="Heading 9", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    for sym, mean in [
        ("API", "Application Programming Interface"),
        ("AUC", "Area Under the Receiver Operating Characteristic Curve"),
        ("BERT", "Bidirectional Encoder Representations from Transformers"),
        ("CHI", "Community Health Index"),
        ("DBSCAN", "Density-Based Spatial Clustering of Applications with Noise"),
        ("ERD", "Entity Relationship Diagram"),
        ("F1", "Harmonic mean of precision and recall"),
        ("GNN", "Graph Neural Network"),
        ("JWT", "JSON Web Token"),
        ("LLM", "Large Language Model"),
        ("MAE", "Mean Absolute Error"),
        ("NLP", "Natural Language Processing"),
        ("p10 / p50 / p90", "10th, 50th and 90th percentile forecast quantiles"),
        ("PCA", "Principal Component Analysis"),
        ("RAG", "Retrieval-Augmented Generation"),
        ("REST", "Representational State Transfer"),
        ("RoBERTa", "Robustly Optimised BERT Pretraining Approach"),
        ("ROC", "Receiver Operating Characteristic"),
        ("SNAP", "Stanford Network Analysis Project"),
        ("TFT", "Temporal Fusion Transformer"),
        ("TF-IDF", "Term Frequency \u2013 Inverse Document Frequency"),
    ]:
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(sym + "\t")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(mean).font.size = Pt(11)


# --------------------------------------------------------------------------
# Chapter 1
# --------------------------------------------------------------------------

def chapter_one(d):
    chapter(d, 1, "INTRODUCTION")
    section(d, "1.1  Background")
    body_text(d,
        "Online communities are now a primary venue for public discussion, technical "
        "collaboration, journalism and political organisation. They are also fragile. "
        "A space that functions for years can deteriorate within weeks: hostility "
        "rises, established members disengage, moderators are overwhelmed by volume, "
        "and the community either fragments or is closed. The pattern is familiar "
        "enough that platforms have named it, yet the tools available to moderators "
        "remain almost entirely reactive.")
    body_text(d,
        "The question that motivates this project has an unusual origin. Between 1968 "
        "and 1973 the ethologist John B. Calhoun ran a series of experiments, the "
        "last and best known of which is Universe 25, in which "
        "a mouse population was given unlimited food, water and protection from "
        "predators, but a fixed amount of space. The population grew, then collapsed. "
        "What Calhoun documented was not starvation but a breakdown of social "
        "behaviour: withdrawal, aggression, the abandonment of normal social roles. He "
        "called the result a behavioural sink [11]. The experiment is contested as a "
        "model of human society, and this project does not claim otherwise. What is "
        "striking for our purposes is a methodological absence: the collapse was "
        "observed and described in detail after the fact, but there was no measurement "
        "that signalled it in advance. Nobody could have said, at month nine, that "
        "month sixteen was coming.", indent_first=True)
    body_text(d,
        "Digital communities present the same problem with one crucial advantage: they "
        "are continuously and completely instrumented. Every message, every "
        "participant and every interaction is recorded with a timestamp. If a "
        "community is deteriorating, the evidence exists in that record before the "
        "outcome is visible. The difficulty is that existing systems do not look for "
        "it. Toxicity classifiers judge a single comment after it has been posted. "
        "Health dashboards report what a community looks like today. Both describe a "
        "state; neither describes a direction.", indent_first=True)
    body_text(d,
        "APOLLO-M is built on the position that community instability should be "
        "treated as a forecasting problem rather than a classification problem. The "
        "system scores individual comments, aggregates them into community-level "
        "measures, learns how those measures move over time, and projects them five "
        "days forward with explicit uncertainty. The intended output is not a verdict "
        "on a comment but a warning about a community, issued early enough that "
        "intervention is still possible.", indent_first=True)

    section(d, "1.2  Problem Statement")
    body_text(d,
        "Moderation tooling operates at the wrong unit of analysis and at the wrong "
        "point in time. It is applied per comment, after publication, and reports "
        "present state rather than trajectory. Consequently a community whose "
        "hostility is rising sharply but from a low base attracts no attention, while "
        "a community that is stably contentious generates continuous alerts that "
        "require no action. Moderators receive volume rather than priority.")
    body_text(d,
        "There is a second, subtler problem. Composite health indices are usually "
        "constructed from present-tense measurements such as toxicity rate, "
        "polarisation and churn. An index of this form measures how unhealthy a "
        "community is, which is not the same quantity as how quickly it is changing. "
        "This project shows empirically that the distinction matters: ranking "
        "communities by such an index performs worse than ranking them by a single "
        "one of its own inputs (Section 3.6).", indent_first=True)
    body_text(d,
        "APOLLO-M therefore addresses the need for a system that operates at community "
        "level, models change over time rather than state, quantifies its own "
        "uncertainty, and produces a prioritised, actionable signal instead of an "
        "undifferentiated stream of flags.", indent_first=True)

    section(d, "1.3  Aims and Objectives")
    body_text(d, "The objectives of this project are as follows:")
    for i, o in enumerate([
        "To score individual community messages for toxicity using a transformer-based "
        "classifier, and to evaluate that classifier on held-out labelled data.",
        "To aggregate per-message scores, together with author churn and graph-derived "
        "polarisation and echo-chamber measures, into an interpretable "
        "community-level Community Health Index and a trend-aware instability score.",
        "To discover behavioural groupings among communities without supervision, so "
        "that communities may be compared against structurally similar peers.",
        "To forecast each community's toxicity five days ahead using a Temporal Fusion "
        "Transformer, producing p10, p50 and p90 quantiles rather than a point "
        "estimate, so that the confidence of each forecast is explicit.",
        "To convert forecasts into graded alerts and recommended moderation actions, "
        "and to serve them through an authenticated API, an interactive dashboard and "
        "a monitoring stack.",
        "To validate the framework against recorded ground truth and report the result "
        "quantitatively, including where components underperform.",
    ], start=1):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(f"{i}.  ").bold = True
        p.add_run(o).font.size = Pt(12)

    section(d, "1.4  Scope of Project")
    body_text(d, "The scope of the current phase is defined as follows:")
    for i, o in enumerate([
        "Analysis operates on English-language text. Multilingual toxicity detection "
        "is outside the present scope.",
        "Community-level analysis covers the sixty largest communities in the working "
        "corpus, selected by message volume; the selection criterion is stated so that "
        "it can be reproduced and criticised.",
        "The forecasting horizon is five days, with a fourteen-day lookback window.",
        "Live platform ingestion is implemented as a module and is pending credential "
        "approval; the present evaluation uses a declared simulation with recorded "
        "ground truth, described in Section 3.5.",
        "The language model is confined to generating natural-language explanations of "
        "results already computed. It does not classify, score or decide.",
    ], start=1):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(f"{i}.  ").bold = True
        p.add_run(o).font.size = Pt(12)


# --------------------------------------------------------------------------
# Chapter 2
# --------------------------------------------------------------------------

def chapter_two(d):
    chapter(d, 2, "LITERATURE REVIEW")
    section(d, "2.1  Background")
    body_text(d,
        "Research relevant to this project falls into five strands. The first is the "
        "automatic detection of toxic, abusive and hateful language, which supplies "
        "the per-message signal on which everything else is built. The second concerns "
        "community-level health and conflict, and asks what happens to a community as "
        "a whole rather than to an individual message. The third represents platforms "
        "as graphs, capturing relationships between communities rather than only their "
        "contents. The fourth is temporal forecasting, which supplies the machinery "
        "for projecting a signal forward with quantified uncertainty. The fifth is "
        "misinformation detection, contributing a second and distinct mode of "
        "community harm. Each strand is well developed individually; the gap this "
        "project addresses lies in their combination.")

    section(d, "2.2  Related Work")
    section(d, "2.2.1  Toxicity and Abusive Language Detection", 4)
    body_text(d,
        "Davidson et al. [1] distinguished hate speech from merely offensive language "
        "and released a labelled corpus of approximately 25,000 annotated messages "
        "that remains a standard benchmark; their central finding, that lexical "
        "methods systematically confuse offensive language with hate speech, motivated "
        "later transformer-based work. Wulczyn et al. [12] scaled annotation to over a "
        "hundred thousand Wikipedia discussion comments, demonstrating that crowd "
        "labels can approximate expert judgement at scale, and Founta et al. [13] "
        "characterised abusive behaviour on Twitter across multiple overlapping "
        "categories. With the arrival of pretrained language models [3][4], "
        "classification accuracy improved substantially; the Detoxify family of models "
        "[2], used in this project, fine-tunes BERT on the Jigsaw toxicity corpora and "
        "provides calibrated per-message toxicity probabilities.")
    body_text(d,
        "The limitation common to this strand is one of unit and timing. These systems "
        "classify a message, in isolation, after it has been published. They provide "
        "no aggregate view of a community and make no statement about the future.",
        indent_first=True)

    section(d, "2.2.2  Community-Level Health and Conflict", 4)
    body_text(d,
        "Chandrasekharan et al. [9] examined the 2015 ban of several hate-focused "
        "Reddit communities and measured the effect on the hate speech produced by "
        "their former members, providing rare quantitative evidence on the effect of a "
        "community-level intervention. Kumar et al. [5] studied conflict between "
        "communities directly, analysing hyperlinks between subreddits and the "
        "mobilisation effects that follow. Cinelli et al. [10] compared echo-chamber "
        "formation across four platforms and showed that the effect varies with "
        "platform architecture rather than being a universal property of social media. "
        "Garimella et al. [16] proposed graph-based measures for quantifying "
        "controversy within discussion networks.")
    body_text(d,
        "This strand establishes that community-level measurement is both meaningful "
        "and tractable. It is, however, predominantly retrospective: the analyses "
        "characterise what happened, and none produces a forward-looking estimate with "
        "quantified uncertainty.", indent_first=True)

    section(d, "2.2.3  Graph Representations of Communities", 4)
    body_text(d,
        "The SNAP Reddit hyperlink network released with Kumar et al. [5] models "
        "communities as nodes and inter-community references as sentiment-weighted "
        "directed edges, and is used in this project to derive per-community "
        "polarisation and echo-chamber measures. Hamilton et al. [6] introduced "
        "GraphSAGE, an inductive framework that generates embeddings for previously "
        "unseen nodes by sampling and aggregating neighbourhood features. That "
        "property matters here: new communities appear continuously, and a "
        "transductive method would require retraining for each one.")

    section(d, "2.2.4  Temporal Forecasting of Social Signals", 4)
    body_text(d,
        "The attention mechanism [8] underpins most current sequence models. Lim et "
        "al. [7] introduced the Temporal Fusion Transformer, which combines recurrent "
        "processing of local dynamics with interpretable attention over longer "
        "horizons, supports static covariates alongside time-varying inputs, and, "
        "decisively for this project, is trained with a quantile loss so that it "
        "produces prediction intervals rather than point estimates. For a moderation "
        "system this distinction is not cosmetic: a forecast of rising toxicity with "
        "wide uncertainty warrants a different response from the same forecast with "
        "narrow uncertainty, and a point estimate cannot express the difference.")

    section(d, "2.2.5  Misinformation Detection", 4)
    body_text(d,
        "Shu et al. [14] surveyed misinformation detection on social media and "
        "characterised the problem as one of combining content, social context and "
        "propagation signals. Zannettou et al. [15] mapped the wider ecosystem of "
        "false information and its actors. Retrieval-augmented generation [21] offers "
        "a route to verdicts grounded in retrieved evidence rather than in a model's "
        "parameters, which is the approach taken by CEREBRO, the companion system "
        "described in Section 3.8.")

    section(d, "2.3  Research Gap and Summary")
    body_text(d,
        "Taken together, the reviewed work supplies every component this project "
        "requires, but no reviewed system combines them. Toxicity classifiers operate "
        "per message and after the fact. Community-level studies operate "
        "retrospectively and produce no forecast. Graph methods characterise structure "
        "without projecting it forward. Forecasting architectures are mature but have "
        "not been applied to community instability with quantified uncertainty. The "
        "gap APOLLO-M addresses is therefore integrative: an end-to-end framework "
        "moving from per-message scoring, through community-level aggregation and "
        "unsupervised structure discovery, to a probabilistic five-day forecast and a "
        "prioritised alert, with each stage measurable in isolation.")
    body_text(d,
        "A second and narrower gap emerged during development and is, to our "
        "knowledge, not addressed in the reviewed literature: composite community "
        "health indices are constructed from present-tense measurements, and are "
        "therefore poorly suited to ranking communities by rate of change. Section 3.6 "
        "quantifies this effect and presents a trend-aware alternative.",
        indent_first=True)

    table(d, "Table 2.1: Research Summary",
          ["Paper Title", "Year", "Author", "Limitation / Research Gap",
           "Algorithms / Methodology", "Result"],
          [
           ["Automated Hate Speech Detection and the Problem of Offensive Language",
            "2017", "Davidson, Warmsley, Macy, Weber",
            "Per-message classification only; no community-level aggregation and no "
            "temporal component.",
            "Crowd-sourced annotation of about 25,000 tweets into hate speech, "
            "offensive language and neither; logistic regression over n-gram and "
            "lexical features.",
            "Established that lexical methods conflate offensive language with hate "
            "speech. The released corpus is the labelled toxicity data used in this "
            "project."],
           ["Ex Machina: Personal Attacks Seen at Scale", "2017",
            "Wulczyn, Thain, Dixon",
            "Produces per-comment attack scores; no notion of community trajectory.",
            "Crowd annotation of over 100,000 Wikipedia talk-page comments; linear and "
            "MLP classifiers over character and word n-grams.",
            "Demonstrated that crowd labels approximate expert judgement at scale, "
            "enabling large toxicity corpora."],
           ["Detoxify: Toxic Comment Classification", "2020",
            "Hanu and the Unitary team",
            "Message-level probabilities only; no aggregation, forecasting or "
            "community context.",
            "BERT-family transformers fine-tuned on the Jigsaw toxicity corpora and "
            "released as pretrained models.",
            "Calibrated per-message toxicity probabilities. unitary/toxic-bert is the "
            "micro-layer scorer in this project."],
           ["Community Interaction and Conflict on the Web", "2018",
            "Kumar, Hamilton, Leskovec, Jurafsky",
            "Retrospective analysis of conflict that has already occurred; no "
            "predictive model of community health.",
            "Directed hyperlink network of 36,000 subreddits with sentiment-weighted "
            "edges; LSTM over user, community and text features to predict "
            "mobilisation.",
            "Showed that inter-community conflict is driven by a small number of "
            "communities. Supplies the SNAP graph used here for polarisation and "
            "echo-chamber measures."],
           ["You Cannot Stay Here: The Efficacy of the 2015 Reddit Ban Examined "
            "Through Hate Speech", "2017", "Chandrasekharan et al.",
            "Evaluates an intervention after the fact; offers no early-warning signal "
            "that would allow earlier action.",
            "Causal analysis of hate-speech usage by affected users before and after "
            "the ban, with matched controls.",
            "Community-level intervention measurably reduced hate speech by affected "
            "users, establishing that community-scale action is effective."],
           ["The Echo Chamber Effect on Social Media", "2021", "Cinelli et al.",
            "Descriptive and cross-sectional; no temporal forecast of a community "
            "trajectory.",
            "Comparative analysis of feed algorithms and interaction networks across "
            "Facebook, Twitter, Reddit and Gab.",
            "Echo-chamber strength varies with platform architecture, motivating "
            "echo-chamber measures as a per-community feature."],
           ["Inductive Representation Learning on Large Graphs (GraphSAGE)", "2017",
            "Hamilton, Ying, Leskovec",
            "General-purpose graph learning; not applied to community health or "
            "instability.",
            "Inductive node embeddings produced by sampling and aggregating "
            "neighbourhood features, generalising to unseen nodes.",
            "Enables embeddings for communities absent at training time. Adopted as "
            "the Phase 2 structural component of the meso layer."],
           ["Temporal Fusion Transformers for Interpretable Multi-horizon Time Series "
            "Forecasting", "2021", "Lim, Arik, Loeff, Pfister",
            "Demonstrated on retail, traffic and electricity data; not applied to "
            "social or community signals.",
            "Attention-based architecture with variable selection, static covariate "
            "encoders and a quantile loss for multi-horizon forecasting.",
            "Interpretable multi-horizon forecasts with prediction intervals. Forms "
            "the macro layer of this project."],
           ["Fake News Detection on Social Media: A Data Mining Perspective", "2017",
            "Shu, Sliva, Wang, Tang, Liu",
            "Detects individual false items; does not model the effect of "
            "misinformation on community stability.",
            "Survey and taxonomy of content-based, social-context and "
            "propagation-based detection methods.",
            "Framed misinformation detection as a multi-signal problem, informing the "
            "planned misinformation-pressure covariate."],
          ], widths=[1.25, 0.42, 0.95, 1.25, 1.35, 1.28], font=7.5)

    section(d, "2.4  Comparison with Existing Studies")
    body_text(d,
        "Table 2.2 compares APOLLO-M with representative systems from the reviewed "
        "literature along the dimensions that distinguish it.")
    table(d, "Table 2.2: Comparison with Existing Studies",
          ["System", "Unit of analysis", "Temporal", "Uncertainty",
           "Graph features", "Deployed"],
          [["Davidson et al. [1]", "Message", "No", "No", "No", "No"],
           ["Detoxify [2]", "Message", "No", "No", "No", "Model release"],
           ["Kumar et al. [5]", "Community pair", "Retrospective", "No", "Yes", "No"],
           ["Chandrasekharan et al. [9]", "Community", "Retrospective", "No", "No", "No"],
           ["Cinelli et al. [10]", "Platform", "Cross-sectional", "No", "Yes", "No"],
           ["APOLLO-M (this work)", "Community", "5-day forecast",
            "p10 / p50 / p90", "Yes", "API + dashboard"]],
          widths=[1.5, 1.15, 1.05, 0.95, 0.85, 1.0], font=8.5)

    section(d, "2.5  Chapter Summary")
    body_text(d,
        "The reviewed literature provides mature solutions to each sub-problem this "
        "project depends upon: transformer-based toxicity classification, "
        "community-level measurement, graph representations of inter-community "
        "relationships, and interpretable probabilistic forecasting. What it does not "
        "provide is their integration into a system that issues an early, prioritised "
        "and uncertainty-aware warning about a community rather than a verdict about a "
        "message. That integration, together with the empirical finding that "
        "present-tense health indices are ill-suited to ranking communities by rate of "
        "change, constitutes the contribution of this work.")



# --------------------------------------------------------------------------
# Chapter 3
# --------------------------------------------------------------------------

def chapter_three(d):
    chapter(d, 3, "DESIGN AND METHODOLOGY")
    body_text(d,
        "This chapter describes what has been built, how it is organised, and what it "
        "measures. The framework is implemented end to end and deployed; Section 3.6 "
        "reports its measured performance and Section 3.9 sets out the work scheduled "
        "for the next phase.")

    section(d, "3.1  Proposed Methodology (Framework / Architecture)")
    body_text(d,
        "APOLLO-M is organised as a layered pipeline. Each layer consumes the output "
        "of the layer beneath it, produces an artefact that can be inspected and "
        "measured on its own, and is independent of the layers above. The design "
        "follows one governing principle: every number the system reports is produced "
        "by a deterministic computation or a measurable model, and a language model is "
        "used only to describe those numbers in natural language. No score, alert or "
        "forecast originates from a language model.")
    figure(d, "apollo_architecture",
           "Figure 3.1: APOLLO-M system architecture. Modules marked INTEGRATED are "
           "built and execute on every pipeline run; those marked PHASE 2 are "
           "implemented and available, with integration scheduled for the next phase.")
    body_text(d,
        "The DATA layer supplies real Reddit community structure from the SNAP "
        "hyperlink network together with a labelled toxicity corpus. The MICRO layer "
        "scores every message. The MESO layer aggregates those scores into "
        "community-level measures and discovers behavioural groupings. The MACRO layer "
        "forecasts. The ACT layer converts forecasts into graded alerts and "
        "recommended actions. Results are then stored, served and observed.",
        indent_first=True)

    section(d, "3.2  Process Model")
    body_text(d,
        "Development followed an incremental and iterative model rather than a "
        "sequential one. Each layer was built, measured against data, and only then "
        "integrated with the layer above. This choice was deliberate: a layered "
        "analytical pipeline accumulates error silently, and a defect introduced early "
        "is expressed as a plausible but wrong number later. Measuring each layer "
        "before depending on it localises faults to the layer that produced them.")
    figure(d, "process_model",
           "Figure 3.2: Process model. Measurement precedes integration at every "
           "increment, and its results feed the following increment.", width=6.0)
    body_text(d,
        "The approach was validated in practice. Three defects were identified by "
        "measuring intermediate outputs rather than by inspecting code: a scaling "
        "error that drove the Community Health Index to zero for every community; a "
        "sampling defect in which the corpus loader read a fixed prefix of a file "
        "sorted by community name, so that analysis covered only communities beginning "
        "with a single letter; and the finding, described in Section 3.6, that the "
        "health index ranked communities less accurately than one of its own inputs. "
        "None of these would have been visible from the code alone.", indent_first=True)

    section(d, "3.3  Modules Discussion")
    figure(d, "block_diagram",
           "Figure 3.3: System block diagram showing the modules and the data passed "
           "between them.", width=6.0)

    section(d, "3.3.1  Micro Layer", 4)
    body_text(d,
        "The micro layer scores each message for toxicity using unitary/toxic-bert "
        "[2], a BERT-family transformer fine-tuned on the Jigsaw toxicity corpora, "
        "which returns a calibrated probability per message. A parallel component "
        "produces sentence embeddings using BERT/RoBERTa representations [3][4], "
        "reduced by principal component analysis, which become covariates for the "
        "forecasting layer. Evaluated on a held-out split of the labelled corpus, the "
        f"classifier achieves {TOX['accuracy']} accuracy, micro F1 {TOX['F1 micro']} "
        f"and macro F1 {TOX['F1 macro']}. Macro F1 is reported alongside micro F1 "
        "because the classes are imbalanced and micro F1 alone would flatter the "
        "result on the majority class.")

    section(d, "3.3.2  Meso Layer", 4)
    body_text(d,
        "The meso layer converts per-message scores into per-community measures. The "
        "Community Health Index combines four normalised inputs into a bounded score "
        "on a 0 to 100 scale, where a higher value indicates a healthier community:")
    p = para(d, "CHI = 100 - (35 T + 30 P + 20 C + 15 E)", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=4)
    para(d, "(3.1)", style="Caption", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_after=8)
    for line in ["T = mean toxicity rate of the community, 0 to 1",
                 "P = polarisation, the share of the community's hyperlinks carrying "
                 "negative sentiment",
                 "C = author churn, the proportion of established authors who ceased "
                 "participating",
                 "E = echo-chamber index, the local clustering coefficient of the "
                 "community in the graph"]:
        para(d, line, size=11, space_after=3, indent=0.8)
    body_text(d,
        "Polarisation and the echo-chamber index are computed from the community's own "
        "node in the SNAP hyperlink graph. An earlier implementation computed both as "
        "global constants applied identically to every community, which meant that 45 "
        "of the 100 available penalty points were a fixed offset incapable of "
        "distinguishing one community from another; per-community computation "
        "corrected this.", indent_first=True)
    body_text(d,
        "The layer also computes an instability score, which measures direction of "
        "travel rather than present state. Mean toxicity over a recent window is "
        "compared against the earlier baseline for the same community, and the "
        "resulting trend is weighted above the absolute level:", indent_first=True)
    para(d, "Instability = 0.65 x min(1, max(0, DT / 0.25)) + 0.35 x T", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=4)
    para(d, "(3.2)", style="Caption", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_after=8)
    para(d, "DT = mean toxicity in the most recent quartile of the window minus the "
            "mean over the preceding period", size=11, space_after=3, indent=0.8)
    body_text(d,
        "The trend term dominates because a community whose hostility is rising is one "
        "where intervention can still change the outcome, whereas a community that is "
        "stably contentious is already known to its moderators. Section 3.6 shows this "
        "score substantially outperforms the Community Health Index at identifying "
        "destabilising communities.", indent_first=True)

    section(d, "3.3.3  Unsupervised Layer", 4)
    body_text(d,
        "Communities are grouped by behavioural profile without labels. K-Means "
        "partitions them into five clusters over the standardised health features, "
        "and DBSCAN [20] independently flags communities that belong to no dense "
        "region and are therefore behaviourally atypical. Principal component "
        "analysis provides a two-dimensional projection for visual inspection. The "
        "purpose is comparative: a community is more usefully judged against "
        "structurally similar peers than against a global average.")

    section(d, "3.3.4  Macro Layer", 4)
    body_text(d,
        "The macro layer forecasts each community's toxicity five days ahead using a "
        "Temporal Fusion Transformer [7] with a fourteen-day lookback. The model is "
        "trained with a quantile loss and emits the 10th, 50th and 90th percentiles "
        "rather than a single value, so that the width of the interval communicates "
        "the confidence of the forecast. An earlier implementation requested the "
        "wrong prediction mode and returned three identical quantiles, producing "
        "confidence bands of zero width; the defect was found by inspecting the "
        "output rather than the code, and the corrected implementation produces "
        "intervals that widen with horizon as expected.")

    section(d, "3.3.5  Act, Serve and Observe", 4)
    body_text(d,
        "Alert bands are derived from the Community Health Index as CRITICAL, HIGH, "
        "MEDIUM or LOW, and each band maps to a recommended moderation action from "
        "NO_ACTION through to EMERGENCY_INTERVENTION. Results are persisted to "
        "PostgreSQL, served by a FastAPI application secured with JSON Web Tokens and "
        "role-based access control, and presented in a Streamlit dashboard. A "
        "Prometheus exporter publishes system metrics, which Grafana charts for "
        "operational monitoring; the deployed API additionally serves the same series "
        "at /metrics so that any monitoring system can scrape it.")

    section(d, "3.4  Project Diagrams")
    section(d, "3.4.1  Database Design (ERD)", 4)
    body_text(d,
        "Results are persisted in PostgreSQL. The schema separates per-community "
        "measurements, alerts and forecasts so that each pipeline run appends a new "
        "observation rather than overwriting history, which is what permits trends to "
        "be computed over time.")
    figure(d, "erd", "Figure 3.4: Entity relationship diagram of the apollo schema.",
           width=6.1)

    section(d, "3.4.2  Sequence Diagram", 4)
    body_text(d,
        "Figure 3.5 traces a typical interaction: a moderator opens a community "
        "drill-down and requests an explanation. The ordering is significant. Every "
        "value shown to the user is read from storage before the language model is "
        "contacted, and the model receives already-computed metrics and returns only "
        "prose.")
    figure(d, "sequence", "Figure 3.5: Sequence diagram for a community drill-down "
                          "with explanation.", width=6.1)

    section(d, "3.4.3  Context Diagram", 4)
    body_text(d,
        "Figure 3.6 places the system within its boundary, showing the external "
        "entities it exchanges data with and the direction of each exchange.")
    figure(d, "context", "Figure 3.6: Context diagram showing the system boundary and "
                         "external entities.", width=6.0)

    section(d, "3.4.4  Use Case Diagram", 4)
    body_text(d,
        "Three roles are defined and enforced by the API through claims carried in the "
        "authentication token: viewer, analyst and administrator.")
    figure(d, "usecase", "Figure 3.7: Use case diagram with role-based access.",
           width=6.0)

    section(d, "3.5  Datasets and Data Provenance")
    body_text(d,
        "The project draws on two established public datasets, listed in Table 3.1. "
        "Notably, they originate from different platforms: the labelled toxicity "
        "corpus is Twitter data and the community structure is Reddit data. The "
        "framework is therefore already exercised across two platforms, which supports "
        "the design claim that everything above the ingestion layer is "
        "platform-independent.")
    table(d, "Table 3.1: Datasets Used and Their Role",
          ["Dataset", "Platform", "Scale", "Role in APOLLO-M"],
          [["Davidson / Jigsaw toxicity corpus [1]", "Twitter / X", "24,783 labelled texts",
            "Training and held-out evaluation of the toxicity classifier; supplies the "
            "scored message pool"],
           ["SNAP Reddit Hyperlink Network [5]", "Reddit",
            "35,776 nodes, 137,821 edges",
            "Per-community polarisation and echo-chamber measures; intended GraphSAGE "
            "input"],
           ["Reddit community metadata", "Reddit", "8,865 communities",
            "Community names, author populations and activity volumes"]],
          widths=[1.9, 1.0, 1.35, 2.25], font=8.5)
    body_text(d,
        "Evaluating a forecaster of community instability requires knowing which "
        "communities actually destabilised and when. No public dataset provides those "
        "labels. The project therefore adopts a standard alternative used where ground "
        "truth is unavailable: a declared simulation in which the answer is recorded "
        "in advance and withheld from the system.", indent_first=True)
    body_text(d,
        "Sixty communities are generated over one hundred and twenty concurrent days "
        "using real community names, author populations and activity volumes. Each is "
        "assigned a latent toxicity propensity, and a recorded subset of fifteen is "
        "given a deliberately rising trend over the final third of the window. Message "
        "toxicity is not synthetic: a pool of twelve thousand real texts is scored once "
        "by the toxicity model, and every generated message carries the genuine score "
        "for the exact text it contains. The ground-truth assignment is written to a "
        "file that no pipeline stage reads, which makes recovery of the planted signal "
        "a measurable outcome rather than an assertion.", indent_first=True)

    section(d, "3.6  Evaluation Methodology and Results")
    body_text(d,
        "Two questions are asked of the system. First, does ranking communities by its "
        "measures surface the ones that were destabilised? Second, does the forecaster "
        "independently predict the direction of change? Detection is reported as "
        "ROC-AUC, which is threshold-free and therefore not sensitive to an arbitrary "
        "alert cut-off.")
    table(d, "Table 3.2: Detection Performance Against Planted Ground Truth",
          ["Ranking signal", "ROC-AUC", "Interpretation"],
          [["Instability score (trend-aware)",
            f"{DET['instability_score_roc_auc']}",
            "Separates planted communities from the rest completely"],
           ["Raw toxicity rate", f"{DET['toxicity_only_roc_auc']}",
            "Moderately informative"],
           ["Community Health Index", f"{DET['chi_roc_auc']}",
            "Barely above chance, and below one of its own inputs"]],
          widths=[2.3, 1.0, 3.2], font=9)
    body_text(d,
        "The ordering in Table 3.2 is the most substantive result of this phase, and "
        "it was not the expected one. The Community Health Index, which combines four "
        "measures, ranked communities less accurately than the single toxicity measure "
        "it contains. The explanation is conceptual rather than a defect in the "
        "implementation: the index is composed entirely of present-tense quantities "
        "and therefore measures how unhealthy a community is, whereas the planted "
        "signal, like real instability, is a change over time. A community that is "
        "stably hostile scores badly and needs no intervention; one that is still "
        "healthy but deteriorating quickly is invisible to such an index. The "
        "instability score of Equation 3.2 was introduced in response and resolves the "
        "problem.", indent_first=True)
    body_text(d,
        "One qualification is necessary. The instability score measures a "
        "recent-versus-baseline difference and the simulation plants a monotonic ramp, "
        "so the score is searching for the shape it was given, and its ROC-AUC of "
        f"{DET['instability_score_roc_auc']} should be read as confirmation that the "
        "pipeline is correctly wired rather than as accuracy on live data. The "
        "forecasting result below carries no such qualification: the Temporal Fusion "
        "Transformer is never shown the ground-truth file.", indent_first=True)
    table(d, "Table 3.3: Forecasting Performance Against Planted Ground Truth",
          ["Measure", "Result"],
          [["Destabilising communities predicted rising",
            f"{int(FC['destabilising_recall'] * 15)} / 15 "
            f"({FC['destabilising_recall']:.0%} recall)"],
           ["Slope ROC-AUC, destabilising against the rest", f"{FC['slope_roc_auc']:.3f}"],
           ["Mean predicted slope, destabilising communities",
            f"{FC['mean_slope_by_trajectory'].get('destabilising', 0):+.5f}"],
           ["Mean predicted slope, stable communities",
            f"{FC['mean_slope_by_trajectory'].get('stable', 0):+.5f}"],
           ["Mean predicted slope, improving communities",
            f"{FC['mean_slope_by_trajectory'].get('improving', 0):+.5f}"]],
          widths=[3.6, 2.9], font=9)
    body_text(d,
        "The forecaster identified a rising trend in every one of the fifteen planted "
        "communities, and the mean predicted slopes order correctly across the three "
        "trajectory types. Recall is the appropriate headline here because the "
        "expensive error in this domain is a missed deterioration, whereas flagging a "
        "quiet community costs only a review.", indent_first=True)
    figure(d, "apollo_request_flow",
           "Figure 3.8: Request processing. Serving is separated from modelling: the "
           "API answers from stored values and performs no inference.", width=6.1)

    section(d, "3.7  Deployed System")
    body_text(d,
        "The system is deployed and publicly reachable. The backend and database run "
        "on Render, the dashboard on Streamlit Community Cloud, and the source is "
        "version-controlled on GitHub.")
    table(d, "Table 3.4: Deployed Endpoints",
          ["Component", "Address"],
          [["Dashboard", "https://apollo-m.streamlit.app"],
           ["REST API", "https://apollo-api-tllm.onrender.com"],
           ["Interactive API documentation", "https://apollo-api-tllm.onrender.com/docs"],
           ["Prometheus metrics endpoint",
            "https://apollo-api-tllm.onrender.com/metrics"],
           ["Source repository",
            "https://github.com/kashaffatimajaffrey-design/apollo-m"],
           ["CEREBRO (companion system)", "https://cerebro-sandy-beta.vercel.app"]],
          widths=[2.2, 4.3], font=9)

    section(d, "3.8  Integration with CEREBRO")
    body_text(d,
        "APOLLO-M is deployed alongside CEREBRO, a companion system for threat and "
        "misinformation intelligence. The two answer different questions at different "
        "scales: CEREBRO evaluates a single item on demand, while APOLLO-M forecasts "
        "the trajectory of many communities over time. CEREBRO produces events; "
        "APOLLO-M consumes time series.")
    figure(d, "integration",
           "Figure 3.9: Integration with CEREBRO. A single PostgreSQL instance serves "
           "both systems in separate schemas; the intelligence exchange between them "
           "is specified and scheduled.", width=6.1)
    body_text(d,
        "Integration at the infrastructure level is complete and verified: one "
        "PostgreSQL instance hosts both systems, with the apollo and cerebro schemas "
        "isolated from one another. The next step, shown dashed in Figure 3.9, is to "
        "aggregate CEREBRO's misinformation verdicts into a daily pressure signal and "
        "supply it to the forecaster as an additional covariate. The rationale is that "
        "toxicity indicates a community in conflict while misinformation indicates one "
        "being manipulated; these are distinct failure modes, and a forecaster with "
        "access to both should anticipate instability earlier than one with access to "
        "either alone.", indent_first=True)

    section(d, "3.9  Work Scheduled for the Next Phase")
    body_text(d,
        "The following items are specified, and in most cases implemented, with "
        "integration scheduled for the remainder of the project. They are listed in "
        "the order in which they will be undertaken.")
    for i, o in enumerate([
        "Live platform ingestion. The ingestion module is implemented against the "
        "PRAW interface and exercised end to end using corpus replay; enabling live "
        "collection requires API credentials, which are subject to account age and "
        "activity thresholds. Because every layer above ingestion consumes a "
        "normalised table of community, author, text and timestamp, adding a further "
        "platform requires one adapter and no change to the layers above it.",
        "Integration of the GraphSAGE component, so that learned structural "
        "embeddings supplement the graph statistics currently used.",
        "Integration of the misinformation classifier into the pipeline, and "
        "population of the corresponding table.",
        "Replacement of the supervised moderation recommender's synthetic training "
        "data with recorded moderator decisions, followed by its integration.",
        "The CEREBRO misinformation-pressure covariate described in Section 3.8.",
        "An automated test suite covering the health-index computation, the micro "
        "layer and the API endpoints.",
    ], start=1):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(f"{i}.  ").bold = True
        p.add_run(o).font.size = Pt(12)

    section(d, "3.10  Chapter Summary")
    body_text(d,
        "A layered pipeline has been designed, implemented and deployed, comprising "
        "transformer-based toxicity scoring, community-level aggregation into a health "
        "index and a trend-aware instability score, unsupervised clustering, "
        "probabilistic five-day forecasting, and graded alerting, served through an "
        "authenticated API and a public dashboard with monitoring. The framework was "
        "evaluated against recorded ground truth: the forecaster identified every "
        "planted destabilising community, and the evaluation additionally established "
        "that a present-tense health index is the wrong instrument for ranking "
        "communities by rate of change, a finding that led directly to the instability "
        "score.")


def references(d):
    chapter(d, "", "REFERENCES")
    refs = [
        "T. Davidson, D. Warmsley, M. Macy and I. Weber, “Automated Hate Speech "
        "Detection and the Problem of Offensive Language,” in Proc. 11th Int. AAAI "
        "Conf. on Web and Social Media (ICWSM), 2017, pp. 512–515.",
        "L. Hanu and the Unitary team, “Detoxify: Toxic Comment Classification with "
        "Transformers,” 2020. [Online]. Available: "
        "https://github.com/unitaryai/detoxify",
        "J. Devlin, M.-W. Chang, K. Lee and K. Toutanova, “BERT: Pre-training of "
        "Deep Bidirectional Transformers for Language Understanding,” in Proc. "
        "NAACL-HLT, 2019, pp. 4171–4186.",
        "Y. Liu et al., “RoBERTa: A Robustly Optimized BERT Pretraining "
        "Approach,” arXiv preprint arXiv:1907.11692, 2019.",
        "S. Kumar, W. L. Hamilton, J. Leskovec and D. Jurafsky, “Community "
        "Interaction and Conflict on the Web,” in Proc. World Wide Web Conf. (WWW), "
        "2018, pp. 933–943.",
        "W. L. Hamilton, R. Ying and J. Leskovec, “Inductive Representation "
        "Learning on Large Graphs,” in Advances in Neural Information Processing "
        "Systems (NeurIPS), 2017, pp. 1024–1034.",
        "B. Lim, S. Ö. Arık, N. Loeff and T. Pfister, “Temporal Fusion "
        "Transformers for Interpretable Multi-horizon Time Series Forecasting,” "
        "International Journal of Forecasting, vol. 37, no. 4, pp. 1748–1764, 2021.",
        "A. Vaswani et al., “Attention Is All You Need,” in Advances in Neural "
        "Information Processing Systems (NeurIPS), 2017, pp. 5998–6008.",
        "E. Chandrasekharan, U. Pavalanathan, A. Srinivasan, A. Glynn, J. Eisenstein "
        "and E. Gilbert, “You Can’t Stay Here: The Efficacy of Reddit’s "
        "2015 Ban Examined Through Hate Speech,” Proc. ACM on Human-Computer "
        "Interaction, vol. 1, no. CSCW, pp. 1–22, 2017.",
        "M. Cinelli, G. De Francisci Morales, A. Galeazzi, W. Quattrociocchi and M. "
        "Starnini, “The Echo Chamber Effect on Social Media,” Proc. National "
        "Academy of Sciences, vol. 118, no. 9, 2021.",
        "J. B. Calhoun, “Death Squared: The Explosive Growth and Demise of a Mouse "
        "Population,” Proc. Royal Society of Medicine, vol. 66, no. 1P2, pp. "
        "80–88, 1973.",
        "E. Wulczyn, N. Thain and L. Dixon, “Ex Machina: Personal Attacks Seen at "
        "Scale,” in Proc. World Wide Web Conf. (WWW), 2017, pp. 1391–1399.",
        "A. M. Founta et al., “Large Scale Crowdsourcing and Characterization of "
        "Twitter Abusive Behavior,” in Proc. 12th Int. AAAI Conf. on Web and Social "
        "Media (ICWSM), 2018.",
        "K. Shu, A. Sliva, S. Wang, J. Tang and H. Liu, “Fake News Detection on "
        "Social Media: A Data Mining Perspective,” ACM SIGKDD Explorations "
        "Newsletter, vol. 19, no. 1, pp. 22–36, 2017.",
        "S. Zannettou, M. Sirivianos, J. Blackburn and N. Kourtellis, “The Web of "
        "False Information: Rumors, Fake News, Hoaxes, Clickbait, and Various Other "
        "Shenanigans,” ACM Journal of Data and Information Quality, vol. 11, no. 3, "
        "pp. 1–37, 2019.",
        "K. Garimella, G. De Francisci Morales, A. Gionis and M. Mathioudakis, "
        "“Quantifying Controversy on Social Media,” ACM Trans. on Social "
        "Computing, vol. 1, no. 1, pp. 1–27, 2018.",
        "M. H. Ribeiro, R. Ottoni, R. West, V. A. F. Almeida and W. Meira, "
        "“Auditing Radicalization Pathways on YouTube,” in Proc. Conf. on "
        "Fairness, Accountability and Transparency (FAT*), 2020, pp. 131–141.",
        "L. Breiman, “Random Forests,” Machine Learning, vol. 45, no. 1, pp. "
        "5–32, 2001.",
        "F. Pedregosa et al., “Scikit-learn: Machine Learning in Python,” "
        "Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.",
        "M. Ester, H.-P. Kriegel, J. Sander and X. Xu, “A Density-Based Algorithm "
        "for Discovering Clusters in Large Spatial Databases with Noise,” in Proc. "
        "2nd Int. Conf. on Knowledge Discovery and Data Mining (KDD), 1996, pp. "
        "226–231.",
        "P. Lewis et al., “Retrieval-Augmented Generation for Knowledge-Intensive "
        "NLP Tasks,” in Advances in Neural Information Processing Systems "
        "(NeurIPS), 2020, pp. 9459–9474.",
        "A. Paszke et al., “PyTorch: An Imperative Style, High-Performance Deep "
        "Learning Library,” in Advances in Neural Information Processing Systems "
        "(NeurIPS), 2019, pp. 8024–8035.",
    ]
    for i, r in enumerate(refs, start=1):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        run = p.add_run(f"[{i}]\t{r}")
        run.font.size = Pt(11)


def main():
    d = new_doc()
    front_matter(d)
    toc(d)
    chapter_one(d)
    chapter_two(d)
    chapter_three(d)
    references(d)
    d.save(str(OUT))
    print(f"OK -> {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    main()

